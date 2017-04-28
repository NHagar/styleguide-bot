# -*- coding: UTF-8 -*-
from docx import *
import pickle
from time import sleep

document = Document("PacificStandardStyleGuide.docx")

terms = []
entry = []
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.bold and ':' in run.text:
            terms.append(entry)
            entry = []
        if len(run.text) > 1:
            entry.append(run.text)
terms = terms[1:-38]
terms[-1] = terms[-1][0:-1]
text_dict = {i[0].strip().lower()[:-1]: ''.join(i[1:]) for i in terms}
text_dict['superlative claims'] = 'Lowercase and use quotes (the “clam capital of the world”)'
text_dict['highways'] = 'Highways are written as follows: U.S. 66, Interstate 90, Illinois 34, Governor Thomas E. Dewey Thruway, Ohio Turnpike, Tri-State Tollway, etc.'
text_dict['fourth of july'] = 'Two caps. Do  not  write 4th of July.'
text_dict['navajo'] = 'Not Navaho'
text_dict['navel-gazing'] = 'Hyphenate per Webster’s'
text_dict['royal'] = 'Lowercase unless it is part of a proper name (Royal Air Force, royal colony, royal family, royal household, Royal Institute of British Architects, Royal Navy, the Royal Society).'
text_dict['city names'] = 'In general give the location of any city on first mention. For U.S. cities identify the state; for foreign cities identify the country (He was born in El Paso, Texas. The museum is located in Glasgow, Scotland.).The following major cities  do not  need to be located: Boston, Chicago, Jerusalem, London, Los Angeles, Mexico City, Moscow, New York City, Paris, Beijing, Rome, San Francisco, Seattle, Shanghai, Tokyo. (However, do give the location for Moscow, Idaho; Paris, Texas; etc.) city inhabitants  (a sampling): Athens (Athenian), Boston (Bostonian), Dublin (Dubliner), Florence (Florentine), Frankfurt (Frankfurter), Hamburg (Hamburger), London (Londoner), Madrid (Madrilenian), Paris (Parisian), Rome (Roman), Stockholm (Stockholmer), Sydney (Sydneysider), Vancouver (Vancouverite), Venice (Venetian), Vienna (Viennese)'
text_dict['civil rights movement'] = 'all lowercase, no hyphen civil titles  (per CMOS 8.21, 8.63)'
text_dict['decades'] = 'Use figures, including the figures for the century, with no apostrophe preferable:  the 1930s , the ’30s (not the thirties).'
text_dict['decision-making'] = 'always hyphenated'
text_dict['her, she'] = 'Do not use these pronouns in referring to countries, storms, ships, or heavenly bodies, except in quoted material. Use  it/its.'
text_dict['greater'] = 'Capitalize when used to designate a community and its surrounding region (Greater Chicago, Greater New York, Greater London).'
text_dict['affect'] = 'Means  influence  (“The moon affects lovers.”). Generally a verb; rarely if ever a noun. See  effect  below.'
text_dict['a.d.'] = 'Use C.E. (/B.C.E.) instead.'
text_dict['company and corporation names'] = 'In ordinary text, generally capitalize and spell out the words  Company, Corporation, Incorporated, Limited,  and  Brothers  when used with the full name of business or industrial firms (General Electric Company; General Motors Corporation). But if a firm uses an abbreviation in its official name or does not set off Incorporated or Limited with a comma, follow the firm’s style.'
text_dict['child care/childcare'] = 'Two words as noun, one word as adjective (childcare service)'
text_dict['nba'] = 'instead of National Basketball Association'
text_dict['indian'] = 'Use Native American or, when possible, the specific names of peoples. For Canadian indigenous peoples, use  First Nation'
text_dict['ellipsis'] = 'When trailing off a sentence, ellipsis followed by period: .... Within a paragraph: “It was the end of an era. ... Three years later, the country rallied.”'
text_dict['new world'] = 'Western Hemisphere; two caps. Do not hyphenate when used adjectively (New World manners).'
text_dict['nfl'] = 'instead of National Football League'
text_dict['nhl'] = 'instead of National Hockey League'
text_dict['farther, further'] = 'Farther  refers to physical distance,  further  to an extension of time or degree.'
text_dict['military titles'] = 'see  Chicago  8.111'
text_dict['aborigines, aboriginal'] = 'Lowercase, except for Australian Aboriginals, then it should be Aboriginal or Aborigine.'
text_dict['powwow'] = 'Avoid. Primary meanings are exclusive to Native Americans; English has co-opted.'
text_dict['west indies'] = 'Do not use this term unless it is in the name of an institution or organization. Use  Caribbean islands  instead.'
text_dict['quoted text'] = 'Use quotation marks and roman type. The quote should follow a colon if having multiple sentences, as in: “She turned to me and said: "I went. Everything exploded."" Contrast with: “She turned to me and said, "I went.""'
text_dict['quran'] = 'no apostrophe'
text_dict['disabilities'] = 'Never use  crippled ,  victims  (AIDS victim),  confined to a wheelchair ,  wheelchair-bound ,  the disabled  the blind/deaf , and  the handicapped . Avoid euphemisms such as  differently abled  or  physically challenged . Use  person/people with a disability  rather than  disabled people . Instead of  handicapped parking/seating  use  accessible seating  or  parking for disabled people .  For specific terms, refer to the   and consult sources to see if they prefer certain language/descriptors.'
text_dict['dr.'] = 'Use only on first reference, for medical doctors.'
text_dict['atm'] = 'Abbreviation is acceptable on first mention.'
text_dict['cafe, cafes'] = 'Accent only if used in venue name.'
text_dict['exotic'] = 'Careful use of this word is advised. While it means  not native, foreign, unusual , when applied to human beings,  exotic  can be ethnocentric and racist (it defines people of color as they relate to whites), or relate to striptease ( exotic dancer ).'
text_dict['primitive'] = 'Use with care; do not use to describe people or their country.'
text_dict['airplane types and names'] = 'Roman, no quotes. But italicize nicknames of specific airplanes.'
text_dict['letters as letters'] = 'italicize per CMOS'
text_dict['life jacket'] = 'two words'
text_dict['best informed, best known, etc.'] = 'Hyphenate such terms when they precede the noun they modify (the best-informed writer; his best-known book; her best-loved poem). Do not hyphenate the terms if they follow the verb (the writer who is best informed; the book for which he is best known).'
text_dict['bible'] = 'Uppercase when referring to the specific religious text (the Old/New Testament); lowercase when referring to a general authoritative text, e.g. “ Elements of Style  is a grammarian’s bible.” '
text_dict['saint'] = 'Abbreviate unless part of official name or name of a particular saint. In French names: no period, hyphenate (St-Rémy).'
text_dict['junior'] = 'For personal names, abbreviate, capitalize, do not set off by comma.'
text_dict['judgment'] = 'judgment > judgement'
text_dict['romantic'] = 'When referring to literature, art, anything inspired by the Romantic movement.'
text_dict['c.e.'] = '(Common Era) and B.C.E. (Before Common Era), caps and periods, no space. In dates, C.E. and B.C.E. follow the year'
text_dict['temperatures'] = 'Use ° and indicate F or C (e.g., 72° F, -12° C). In giving ranges, use the degree symbol with both readings but the C or F only with the second reading:  from 68° to 95° F.  In the case of a rough estimate, format as follows:  Temperatures often stay in the high 80s F.  (No need to use “F” or “C” once Fahrenheit or Celsius has been established.) Per Chi, the degree symbol goes with the F or C, and the whole thing is set solid—68°F. Whether you put a space in there or not, the ° symbol  always  goes with the F or C  not  with the numeral. My preferred form—76 °F (with space after numeral).'
text_dict['-theme'] = 'Do not use “-themed.”'
text_dict['20-somethings'] = 'not twenty-somethings; but preferably avoid this phrase'
text_dict['government/government bodies'] = 'Lowercase the word even if customarily used to refer to the cabinet and prime minister of a country. Capitalize the names of specific government agencies, departments, and other units. In most cases, such generic terms as department, agency, bureau, and commission are lowercased if they stand alone in second references.'
text_dict['gray'] = 'not grey'
text_dict['alternate spellings'] = 'When m-w.com lists alternate spellings of a word, use the first one listed unless listed as exception in style guide'
text_dict['crowd surfing'] = 'Two words, no hyphen'
text_dict['cultural movements and styles'] = 'See  Chicago  8.85 for more. Search on “rococo” in Frommers.com for the distinction between baroque and rococo. art deco art nouveau baroque beaux arts classical, classicism dadaism, dada deconstruction Doric Epicurean existentialism Gothic humanism impressionism modernism neoclassical, neoclassicism op art pop art postmodernism rococo romantic, romanticism surrealism'
text_dict['numbers'] = 'Follow these guidelines: Spell out  one  through  nine ; use numerals for  10  and up. Spell out twenties, thirties, forties, etc.  Spell out  one million ,  two million ;  10 million  and up. Same for billions. Use numerals for money: $1 million, $11 million; $1 billion, $2 billion.... Use numerals for percentages: 4 percent, 12 percent. Spell out centuries up to ninth; use numerals for 10th and up. Apply above rules even if a number less than 10 appears in the same phrase as one above 10 (“between four and 11 days”). Always spell out numbers that begin a sentence. Numbers of four digits or more get commas except in years and addresses. Spell out ordinals up to ninth; use numerals for 10th and up (e.g., twenty-secondth would be 22th; three-thousand-seventy-seventh would be 3,077th). Use superscripts with ordinals.'
text_dict['party'] = 'Usually lowercase, even in most party names—Democratic party, Conservative party. Cap Communist Party when it is the only party in the nation.'
text_dict['percent'] = 'Use hyphen in compounds (“10-percent discount”).'
text_dict['magazine department names'] = 'Caps and small caps in editorial text if self referential.'
text_dict['-maker, -making'] = 'Set solid as suffix, no hyphen'
text_dict['minority'] = 'Do not use when referring to an ethnic group'
text_dict['mlb'] = 'instead of Major League Baseball'
text_dict['arctic'] = 'the Arctic, Arctic Circle, arctic climate, Arctic waters, North Pole, the Pole, North Polar ice cap, polar region, polar climate.'
text_dict['religious groups'] = 'Christian, Muslim, Jewish, Buddhist…'
text_dict['coast'] = 'Capitalize the word only when referring to specific geographic regions of the United States lying along specific coasts (Atlantic Coast, Gulf Coast, Pacific Coast, East Coast, West Coast). Lowercase coast in subsequent references if standing alone.'
text_dict['co-exist'] = 'hyphenated'
text_dict['plos one'] = 'Not “PLOS ONE” (the journal’s 2012, all-caps logo rebranding)'
text_dict['policymakers'] = 'always one word'
text_dict['age'] = 'Use “and older” and “and younger,” not over, under, or above. She was two years old. They were in their 90s.'
text_dict['well'] = 'Hyphenate compounds with  well  only when they precede the noun modified (well-known author, well-organized program). Note, however, that  well  compounds are not hyphenated when they precede a noun if the expression itself has a modifier (an exceptionally well known author, a remarkably well organized program). Others: well-dressed, well-being, well-to-do, well wishers.'
text_dict['rv, rvers'] = 'Okay to abbreviate, no periods'
text_dict['sage publications'] = 'Small caps for  sage'
text_dict['nanomaterials'] = 'one word, lowercase'
text_dict['sexism'] = 'Avoid sexist language, such as masculine nouns and pronouns if the subject matter could apply to women as well as to men. Use  –person  in place of  -man  in such words as  businessman,  or recast. Do not use the suffix  “-ess.”'
text_dict['skivvies, skivvies'] = 'Cap for underwear; lowercase for “female domestic servants”'
text_dict['captions'] = 'All caps bold for kickers, no colon.'

doc2 = Document("Styleodds.docx")
terms2 = []
entry2 = []

for paragraph in doc2.paragraphs:
    terms2.append(paragraph.text)

terms2 = [i for i in terms2 if len(i) > 0]
terms2 = [i.split(':') for i in terms2]

text_dict2 = {i[0].strip().lower(): ''.join(i[1:]) for i in terms2}
print len(text_dict)
text_dict.update(text_dict2)
pickle.dump(text_dict, open('style_guide', 'wb'))
print len(text_dict)

import pickle
import unicodedata
_ = pickle.load(open('style_guide', 'rb'))

for i in _.keys():
    if isinstance(i, unicode):
        j = unicodedata.normalize('NFKD', i).encode('ascii','ignore')
        _[j] = _[i]
        del _[i]

for i in _.keys():
    print type(i)
print len(_)
pickle.dump(_, open('style_guide', 'wb'))
